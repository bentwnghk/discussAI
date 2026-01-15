# -*- coding: utf-8 -*-
import concurrent.futures as cf
import glob
import io
import os
import time
import datetime
import base64
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import List, Literal, Optional, Dict, Any
import json
import gradio as gr
import sentry_sdk
from fastapi import FastAPI
from fastapi.staticfiles import StaticFiles
from loguru import logger # logger is imported here
from openai import OpenAI
from promptic import llm
from pydantic import BaseModel, ValidationError
from pypdf import PdfReader
from tenacity import retry, retry_if_exception_type, stop_after_attempt, wait_exponential # tenacity is imported here
from mimetypes import guess_type
import docx # Added for DOCX support
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import pytz
import requests
import html
from bs4 import BeautifulSoup

OPENAI_VOICE_MAPPINGS = {
    "Candidate A": "nova",
    "Candidate B": "alloy",
    "Candidate C": "fable",
    "Candidate D": "echo",
}

TRANSCRIPT_COLORS = {
    "Candidate A": "#E3F2FD",  # Light blue
    "Candidate B": "#FFFDE7",  # Light yellow
    "Candidate C": "#E8F5E8",  # Light green
    "Candidate D": "#FDECEA",  # Light pink
}

if sentry_dsn := os.getenv("SENTRY_DSN"):
    sentry_sdk.init(sentry_dsn)

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")


class DialogueItem(BaseModel):
    text: str
    speaker: Literal["Candidate A", "Candidate B", "Candidate C", "Candidate D"]

    def voice(self): # Remove language parameter, make it a property again
        # Always use OPENAI_VOICE_MAPPINGS for English voices.
        # one-api will be responsible for mapping these voice IDs (e.g., "nova")
        # to the correct downstream provider voice for English.
        return OPENAI_VOICE_MAPPINGS[self.speaker]


class LearningNotes(BaseModel):
    ideas: str  # Structured outline of ideas
    language: str  # Vocabulary
    communication_strategies: str  # Interaction strategies used

class Dialogue(BaseModel):
    scratchpad: str
    dialogue: List[DialogueItem]
    learning_notes: LearningNotes


# Add retry mechanism to TTS calls for resilience
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10), retry=retry_if_exception_type(Exception))
def get_mp3(text: str, voice: str, api_key: str = None) -> bytes:
    """
    Generates MP3 audio for the given text by making a direct request to the
    one-api compatible endpoint (expected at OPENAI_BASE_URL).
    The 'language_selection' parameter is included in the payload to one-api,
    allowing one-api to route to appropriate downstream TTS providers (e.g., MiniMax)
    and apply necessary language boosts. Includes retries.
    """
    effective_api_key = api_key or os.getenv("OPENAI_API_KEY")
    base_url = os.getenv("OPENAI_BASE_URL") # This should be your one-api base URL

    if not effective_api_key:
        logger.error("API key is not configured.")
        raise ValueError("API key not configured.")
    if not base_url:
        logger.error("Base URL is not configured.")
        raise ValueError("Base URL not configured.")

    # Construct the full URL to the one-api speech endpoint
    # Assumes one-api exposes an OpenAI-compatible speech endpoint at /v1/audio/speech
    speech_endpoint_url = f"{base_url.rstrip('/')}/audio/speech"

    headers = {
        "Authorization": f"Bearer {effective_api_key}",
        "Content-Type": "application/json"
    }

    # Manually construct the payload for English TTS
    payload: Dict[str, Any] = { # Ensure payload is explicitly typed for clarity
        "model": "tts-1",
        "voice": voice,
        "input": text,
        "response_format": "mp3",
    }

    logger.debug(
        f"Requesting TTS. Endpoint: '{speech_endpoint_url}', "
        f"Voice: '{voice}', Language: English, Text: '{text[:50]}...'"
    )

    try:
        response = requests.post(
            speech_endpoint_url,
            headers=headers,
            json=payload,
            timeout=60.0 # Timeout in seconds
        )
        response.raise_for_status() # Raise an exception for HTTP error codes (4xx or 5xx)

        logger.debug(
            f"TTS generation successful. Voice: '{voice}', Text: '{text[:50]}...'"
        )
        return response.content  # The binary content of the MP3

    except requests.exceptions.HTTPError as http_err:
        error_message = f"HTTP error during TTS generation. Voice: '{voice}', Text: '{text[:50]}...'. " \
                        f"Status: {http_err.response.status_code if http_err.response else 'N/A'}. " \
                        f"Error: {http_err}. " \
                        f"Response: {http_err.response.text if http_err.response else 'No response text'}"
        logger.error(error_message)
        raise # Reraise exception to trigger tenacity retry
    except requests.exceptions.RequestException as req_err: # Catches other requests errors (e.g., timeout, connection error)
        logger.error(
            f"Request error during TTS generation. Voice: '{voice}', Text: '{text[:50]}...'. Error: {req_err}"
        )
        raise # Reraise exception to trigger tenacity retry
    except Exception as e: # Catch-all for other unexpected errors
        logger.error(
            f"Unexpected error during TTS generation. Voice: '{voice}', Text: '{text[:50]}...'. Error: {e}"
        )
        raise # Reraise exception to trigger tenacity retry

def is_pdf(filename):
    if not filename: return False
    t, _ = guess_type(filename)
    # Check extension and guessed MIME type
    return filename.lower().endswith(".pdf") or (t or "").endswith("pdf")

def is_image(filename):
    if not filename: return False
    t, _ = guess_type(filename)
    image_exts = (".jpg", ".jpeg", ".png")
    # Check extension and guessed MIME type
    return filename.lower().endswith(image_exts) or (t or "").startswith("image")

def is_docx(filename):
    if not filename: return False
    t, _ = guess_type(filename)
    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    # Check extension and guessed MIME type
    return filename.lower().endswith(".docx") or (t or "") == docx_mime

# Add retry mechanism to Vision calls
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10), retry=retry_if_exception_type(Exception))
def extract_text_from_image_via_vision(image_file, openai_api_key=None):
    """Extracts text from an image using OpenAI Vision API, with retries."""
    client = OpenAI(
        api_key=openai_api_key or os.getenv("OPENAI_API_KEY"),
        base_url=os.getenv("OPENAI_BASE_URL"),
        timeout=120.0 # Longer timeout for vision potentially
    )
    logger.debug(f"Requesting Vision text extraction for image: {image_file}")
    try:
        with open(image_file, "rb") as f:
            data = f.read()
            mime_type = guess_type(image_file)[0] or "image/png" # Default to png if guess fails
            b64 = base64.b64encode(data).decode("utf-8")
            image_url = f"data:{mime_type};base64,{b64}"

        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": image_url,
                            "detail": "auto" # Use 'auto' or 'high' based on needs
                        }
                    },
                    {
                        "type": "text",
                        "text": "Extract all the computer-readable text from this image as accurately as possible. Avoid commentary, return only the extracted text."
                    },
                ]
            }
        ]

        response = client.chat.completions.create(
            model="gpt-4.1-mini", # Ensure this model supports vision
            messages=messages,
            max_tokens=32768,
            temperature=0,
        )
        extracted_text = response.choices[0].message.content.strip()
        logger.debug(f"Vision extraction successful for {image_file}. Text length: {len(extracted_text)}")
        return extracted_text
    except Exception as e:
        logger.error(f"Vision extraction failed for {image_file}. Error: {e}")
        raise # Reraise for retry


# Normal mode dialogue generation function
@retry(stop=stop_after_attempt(2), wait=wait_exponential(multiplier=1, min=4, max=15), retry=retry_if_exception_type(ValidationError))
@llm(
    model=os.getenv("OPENAI_MODEL_NORMAL"),
    api_key=os.getenv("OPENAI_API_KEY"),
    base_url=os.getenv("OPENAI_BASE_URL"),
    provider="openai",
    temperature=0.5,
    max_tokens=16384,
    timeout=150.0
)
def generate_dialogue_normal(text: str) -> Dialogue:
    """
    You are an English language tutor helping Hong Kong secondary students improve their speaking skills, especially for group discussions in oral exams.

    Your task is to take the input text provided and create a realistic group discussion in English between four students (Candidate A, B, C, D) on the topic provided in the input text. Don't worry about the formatting issues or any irrelevant information; your goal is to extract the discussion topic and question prompts as well as any relevant key points or interesting facts from the input text for the group discussion.

    Important: The ENTIRE dialogue (including brainstorming, scratchpad, and actual dialogue) should be written in English.

    Here is the input text you will be working with:

    <input_text>
    {text}
    </input_text>

    First, carefully read through the input text and identify the discussion topic and question prompts, as well as any relevant key points or interesting facts from the text accompanying the discussion topic.

    <scratchpad>
    Brainstorm ideas and outline the discussion. Make sure your discussion follows the question prompts you identified in the input text.

    Express a range of well-developed ideas clearly, with elaboration and detail.

    Model an authentic discussion and interaction among 4 students, and include all the following interaction strategies:
    - Strategies for initiating a group discussion (e.g. Alright, we are here to discuss the proposal to ... | Let's begin by talking about the reasons why ...).
    - Strategies for maintaining a group discussion (e.g. What do you think? | Any thoughts, Candidate C?).
    - Strategies for transitioning in a group discussion (e.g. Does anyone have anything else to add? If not, shall we move on and discuss ...?).
    - Strategies for responding in a group discussion (e.g. I agree. | That's an interesting suggestion, but I'm a bit worried that ... | Sorry, I disagree.).
    - Strategies for rephrasing a group discussion (e.g. I see what you mean. You were saying that ...).
    - Strategies for asking for clarification in a group discussion (e.g. I'm not sure if I understood you correctly. Did you mean that ...?).

    Ensure every student contributes their ideas to every question prompt.

    Use natural, accurate vocabulary and expressions suitable for Hong Kong secondary students.

    Write your brainstorming ideas and discussion outline here.
    </scratchpad>

    Now that you have brainstormed ideas and created an outline, it's time to write the full dialogue.

    <podcast_dialogue>
    Write an engaging, informative dialogue here that will be 6-7 minutes long when spoken at a natural pace.

    Use a conversational tone.

    Include all the above interaction strategies to extend the interaction naturally.

    Use 'Candidate A', 'Candidate B', 'Candidate C', 'Candidate D' to identify the 4 speakers. Do not include any bracketed placeholders like [Candidate A] or [Candidate B].

    Alternate speakers naturally, ensuring every candidate speaks 4-6 times throughout the discussion.

    Design your output to be read aloud -- it will be directly converted into audio.

    Assign appropriate speakers (Candidate A, Candidate B, Candidate C, Candidate D) to each line. Ensure the output strictly adheres to the required format: a list of objects, each with 'text' and 'speaker' fields.

    Make the dialogue 6-7 minutes long when spoken at a natural pace (approximately 120-150 words per minute).

    At the end of the dialogue, include a brief summary (1–2 sentences) by one of the candidates.
    </podcast_dialogue>

    <learning_notes>
    Now create comprehensive learning notes for Hong Kong secondary students based on the dialogue you just generated. The learning notes should have three sections:

    **1. Ideas Section:**
    Create a structured outline showing the main ideas discussed in the dialogue. Format this as HTML with proper structure:
    - Use <strong> tags to bold main question prompts or key topics
    - Use <em> tags to italicize important concepts or emphasis
    - Use <br><br> for line breaks between major points
    - Use bullet points (•) or numbered lists with <br> after each item
    - Create clear hierarchy with indentation using &nbsp;&nbsp;&nbsp;&nbsp; for sub-points
    - Reference the question prompts from the input text and show how the discussion addressed each one
    - Include Traditional Chinese translations for all major points and sub-points
    
    Example format:
    <strong>Question 1: [Topic]</strong><br><br>
    • Main point 1<br>
    &nbsp;&nbsp;&nbsp;&nbsp;- Sub-point with <em>emphasis</em><br>
    &nbsp;&nbsp;&nbsp;&nbsp;- Another sub-point<br>
    • Main point 2<br><br>

    **2. Language Section:**
    Identify 12-15 useful vocabulary words from the dialogue. For each item:
    - Provide the English word/phrase
    - Give the Traditional Chinese translation (繁體中文)
    - Show how it was used in the dialogue with a brief example

    Format this as an HTML table with proper structure:
    <table>
    <tr><th>English</th><th>中文</th><th>Usage Example</th></tr>
    <tr><td><strong>word/phrase</strong></td><td>中文翻譯</td><td>Example sentence from dialogue</td></tr>
    </table>

    **3. Communication Strategies Section:**
    List and explain 6-8 interaction strategies that were demonstrated in the dialogue. Format this as HTML with proper structure:
    - Use <strong> tags to bold strategy names
    - Use <em> tags to italicize example phrases from the dialogue
    - Use <br><br> for line breaks between different strategies
    - Use <br> after each example phrase
    - Include Traditional Chinese explanations
    
    Example format:
    <strong>1. Initiating Discussion (開始討論)</strong><br>
    • <em>"Alright, we’re here to discuss [whether our school should serve plant-based meats in the canteen on Green Mondays]."</em><br>
    • <em>"Maybe we can start by talking about [why plant-based meats are becoming so popular these days]."</em><br>
    用於開始討論的策略，幫助引導話題方向。<br><br>

    <strong>2. Maintaining Discussion (維持討論)</strong><br>
    • <em>"What do you think?"</em><br>
    • <em>"Any thoughts on this?"</em><br>
    用於鼓勵其他人參與討論。<br><br>

    Strategies to include:
    - Initiating discussion (開始討論)
    - Maintaining discussion (維持討論)
    - Transitioning between topics (轉換話題)
    - Responding and agreeing/disagreeing (回應及表達同意/不同意)
    - Elaborating with examples (舉例說明)
    - Building on others' ideas (延伸他人想法)
    - Asking for clarification (要求澄清)
    - Rephrasing (重新表述)
    - Summarizing (總結)

    Write all learning notes content in a mix of English and Traditional Chinese to facilitate Hong Kong students' learning.
    </learning_notes>
    """

# Deeper mode dialogue generation function
@retry(stop=stop_after_attempt(2), wait=wait_exponential(multiplier=1, min=4, max=15), retry=retry_if_exception_type(ValidationError))
@llm(
    model=os.getenv("OPENAI_MODEL_DEEP"),
    api_key=os.getenv("OPENAI_API_KEY"),
    base_url=os.getenv("OPENAI_BASE_URL"),
    provider="openai",
    temperature=0.5,
    max_tokens=16384,
    timeout=150.0
)
def generate_dialogue_deeper(text: str) -> Dialogue:
    """
    You are an English language tutor helping Hong Kong secondary students improve their speaking skills, especially for group discussions in oral exams.

    Your task is to take the input text provided and create a realistic group discussion in English between four students (Candidate A, B, C, D) on the topic provided in the input text. Don't worry about the formatting issues or any irrelevant information; your goal is to extract the discussion topic and question prompts as well as any relevant key points or interesting facts from the input text for the group discussion.

    Important: The ENTIRE dialogue (including brainstorming, scratchpad, and actual dialogue) should be written in English.

    Here is the input text you will be working with:

    <input_text>
    {text}
    </input_text>

    First, carefully read through the input text and identify the discussion topic and question prompts, as well as any relevant key points or interesting facts from the text accompanying the discussion topic.

    <scratchpad>
    Brainstorm ideas and outline the discussion. Make sure your discussion follows the question prompts you identified in the input text.

    Express a range of well-developed ideas clearly, with elaboration and detail.

    Model an authentic discussion and interaction among 4 students, and include all the following interaction strategies:
    - Strategies for initiating a group discussion (e.g. Alright, we are here to discuss the proposal to ... | Let's begin by talking about the reasons why ...).
    - Strategies for maintaining a group discussion (e.g. What do you think? | Any thoughts, Candidate C?).
    - Strategies for transitioning in a group discussion (e.g. Does anyone have anything else to add? If not, shall we move on and discuss ...?).
    - Strategies for responding in a group discussion (e.g. I agree. | That's an interesting suggestion, but I'm a bit worried that ... | Sorry, I disagree.).
    - Strategies for rephrasing a group discussion (e.g. I see what you mean. You were saying that ...).
    - Strategies for asking for clarification in a group discussion (e.g. I'm not sure if I understood you correctly. Did you mean that ...?).

    Ensure every student contributes their ideas to every question prompt.

    Use natural, accurate vocabulary and expressions suitable for Hong Kong secondary students.

    Write your brainstorming ideas and discussion outline here.
    </scratchpad>

    Now that you have brainstormed ideas and created an outline, it's time to write the full dialogue.

    <podcast_dialogue>
    Write an engaging, informative dialogue here that will be 6-7 minutes long when spoken at a natural pace.

    Use a conversational tone.

    Include all the above interaction strategies to extend the interaction naturally.

    Use 'Candidate A', 'Candidate B', 'Candidate C', 'Candidate D' to identify the 4 speakers. Do not include any bracketed placeholders like [Candidate A] or [Candidate B].

    Alternate speakers naturally, ensuring every candidate speaks 4-6 times throughout the discussion.

    Design your output to be read aloud -- it will be directly converted into audio.

    Assign appropriate speakers (Candidate A, Candidate B, Candidate C, Candidate D) to each line. Ensure the output strictly adheres to the required format: a list of objects, each with 'text' and 'speaker' fields.

    Make the dialogue 6-7 minutes long when spoken at a natural pace (approximately 120-150 words per minute).

    At the end of the dialogue, include a brief summary (1–2 sentences) by one of the candidates.
    </podcast_dialogue>

    <learning_notes>
    Now create comprehensive learning notes for Hong Kong secondary students based on the dialogue you just generated. The learning notes should have three sections:

    **1. Ideas Section:**
    Create a structured outline showing the main ideas discussed in the dialogue. Format this as HTML with proper structure:
    - Use <strong> tags to bold main question prompts or key topics
    - Use <em> tags to italicize important concepts or emphasis
    - Use <br><br> for line breaks between major points
    - Use bullet points (•) or numbered lists with <br> after each item
    - Create clear hierarchy with indentation using &nbsp;&nbsp;&nbsp;&nbsp; for sub-points
    - Reference the question prompts from the input text and show how the discussion addressed each one
    - Include Traditional Chinese translations for all major points and sub-points
    
    Example format:
    <strong>Question 1: [Topic]</strong><br><br>
    • Main point 1<br>
    &nbsp;&nbsp;&nbsp;&nbsp;- Sub-point with <em>emphasis</em><br>
    &nbsp;&nbsp;&nbsp;&nbsp;- Another sub-point<br>
    • Main point 2<br><br>

    **2. Language Section:**
    Identify 10-15 useful vocabulary words from the dialogue. For each item:
    - Provide the English word/phrase
    - Give the Traditional Chinese translation (繁體中文)
    - Show how it was used in the dialogue with a brief example

    Format this as an HTML table with proper structure:
    <table>
    <tr><th>English</th><th>中文</th><th>Usage Example</th></tr>
    <tr><td><strong>word/phrase</strong></td><td>中文翻譯</td><td>Example sentence from dialogue</td></tr>
    </table>

    **3. Communication Strategies Section:**
    List and explain 6-8 interaction strategies that were demonstrated in the dialogue. Format this as HTML with proper structure:
    - Use <strong> tags to bold strategy names
    - Use <em> tags to italicize example phrases from the dialogue
    - Use <br><br> for line breaks between different strategies
    - Use <br> after each example phrase
    - Include Traditional Chinese explanations
    
    Example format:
    <strong>1. Initiating Discussion (開始討論)</strong><br>
    • <em>"Alright, we’re here to discuss [whether our school should serve plant-based meats in the canteen on Green Mondays]."</em><br>
    • <em>"Maybe we can start by talking about [why plant-based meats are becoming so popular these days]."</em><br>
    用於開始討論的策略，幫助引導話題方向。<br><br>

    <strong>2. Maintaining Discussion (維持討論)</strong><br>
    • <em>"What do you think?"</em><br>
    • <em>"Any thoughts on this?"</em><br>
    用於鼓勵其他人參與討論。<br><br>

    Strategies to include:
    - Initiating discussion (開始討論)
    - Maintaining discussion (維持討論)
    - Transitioning between topics (轉換話題)
    - Responding and agreeing/disagreeing (回應及表達同意/不同意)
    - Elaborating with examples (舉例說明)
    - Building on others' ideas (延伸他人想法)
    - Asking for clarification (要求澄清)
    - Rephrasing (重新表述)
    - Summarizing (總結)

    Write all learning notes content in a mix of English and Traditional Chinese to facilitate Hong Kong students' learning.
    </learning_notes>
    """

def generate_audio(
    input_method: str,
    files: Optional[List[str]],
    input_text: Optional[str],
    dialogue_mode: str = "Normal",
    openai_api_key: str = None,
) -> (str, str, str, str): # Added 4th str for the hidden gr.File component
    """Generates audio from uploaded files or direct text input."""
    start_time = time.time()
    
    # API Key Check - one-api (at OPENAI_BASE_URL via resolved_openai_api_key) handles all TTS routing.
    # It needs its own API key (OPENAI_API_KEY or the one from UI input).
    if not (openai_api_key or os.getenv("OPENAI_API_KEY")): # Check if any source provides the key for one-api
        raise gr.Error("Mr.🆖 AI Hub API Key is required.")

    # Resolve OpenAI API key and Base URL once (used for dialogue and audio generation)
    resolved_openai_api_key = openai_api_key or os.getenv("OPENAI_API_KEY")
    resolved_openai_base_url = os.getenv("OPENAI_BASE_URL")
    
    full_text = ""
    gr.Info("🔍 Analysing group interaction task...")
    podcast_title_base = "Group Discussion" # Default title base

    if input_method == "Upload Files":
        if not files:
            raise gr.Error("Please upload at least one file or switch to another input method.")
        texts = []
        file_names = []
        for file_path in files:
            if not file_path:
                logger.warning("Received an empty file path in the list, skipping.")
                continue
            actual_file_path = file_path.name if hasattr(file_path, 'name') else file_path
            file_path_obj = Path(actual_file_path)
            file_names.append(file_path_obj.stem)
            logger.info(f"Processing file: {file_path_obj.name}")
            text = ""

            if is_pdf(str(file_path_obj)):
                try:
                    with file_path_obj.open("rb") as f:
                        reader = PdfReader(f)
                        if reader.is_encrypted:
                             logger.warning(f"Skipping encrypted PDF: {file_path_obj.name}")
                             raise gr.Error(f"Cannot process password-protected PDF: {file_path_obj.name}")
                        page_texts = [page.extract_text() for page in reader.pages if page.extract_text()]
                        text = "\n\n".join(page_texts) if page_texts else ""
                        if not text: logger.warning(f"No text extracted from PDF: {file_path_obj.name}")
                except Exception as e:
                    logger.error(f"Error reading PDF {file_path_obj.name}: {e}")
                    if "PdfReadError" in str(type(e)):
                         raise gr.Error(f"Error reading PDF file: {file_path_obj.name}. It might be corrupted or improperly formatted.")
                    else:
                         raise gr.Error(f"Error processing PDF file: {file_path_obj.name}.")
            elif is_image(str(file_path_obj)):
                try:
                    text = extract_text_from_image_via_vision(str(file_path_obj), resolved_openai_api_key)
                except Exception as e:
                    logger.error(f"Error processing image {file_path_obj.name} with Vision API: {e}")
                    raise gr.Error(f"Error extracting text from image: {file_path_obj.name}. Check API key, file format, and OpenAI status. Error: {e}")
            elif is_docx(str(file_path_obj)):
                try:
                    doc = docx.Document(actual_file_path)
                    paragraphs = [p.text for p in doc.paragraphs if p.text]
                    text = "\n\n".join(paragraphs)
                    if not text: logger.warning(f"No text extracted from DOCX: {file_path_obj.name}")
                except Exception as e:
                    logger.error(f"Error reading DOCX file {file_path_obj.name}: {e}")
                    if "PackageNotFoundError" in str(type(e)):
                        raise gr.Error(f"Error reading DOCX file: {file_path_obj.name}. It might be corrupted, not a valid DOCX format, or password-protected.")
                    else:
                        raise gr.Error(f"Error processing DOCX file: {file_path_obj.name}.")
            else:
                try:
                   f_size = file_path_obj.stat().st_size
                   if f_size > 0:
                       raise gr.Error(f"Unsupported file type: {file_path_obj.name}. Please upload DOCX, PDF, or image file (JPG, JPEG, PNG). Note: Older .doc format is not supported.")
                   else:
                       logger.warning(f"Skipping empty or placeholder file: {file_path_obj.name}")
                       text = ""
                except FileNotFoundError:
                    logger.warning(f"File not found during processing, likely a temporary file issue: {actual_file_path}")
                    text = ""
                except Exception as e:
                     logger.error(f"Error checking file status for {file_path_obj.name}: {e}")
                     raise gr.Error(f"Error accessing file: {file_path_obj.name}.")
            texts.append(text)
        full_text = "\n\n".join(filter(None, texts))
        if not full_text.strip():
             raise gr.Error("Could not extract any text from the uploaded file(s). Please check the files or try different ones.")
        if file_names:
            podcast_title_base = file_names[0] if len(file_names) == 1 else f"{len(file_names)} Files"


    elif input_method == "Enter Topic":
        if not input_text or not input_text.strip():
            raise gr.Error("Please enter topic or switch to another input method.")
        full_text = input_text
        podcast_title_base = "Pasted Topic"

    

    else:
        raise gr.Error("Invalid input method selected.")

    logger.info(f"Total input text length: {len(full_text)} characters.")
    if not full_text.strip(): # Double check after all input methods
        raise gr.Error("No text content to process. Please provide valid input.")

    # Choose dialogue generation function based on mode
    if dialogue_mode == "Deeper":
        dialogue_generator = generate_dialogue_deeper
    else:
        dialogue_generator = generate_dialogue_normal

    try:
        gr.Info("✨ Generating dialogue and study notes... (This may take 1-2 minutes.)")
        llm_start_time = time.time()
        llm_output = dialogue_generator(full_text)
        logger.info(f"Dialogue generation took {time.time() - llm_start_time:.2f} seconds.")

    except ValidationError as e:
        logger.error(f"LLM output validation failed after retries: {e}")
        raw_output = getattr(e, 'llm_output', str(e)) 
        raise gr.Error(f"The AI model returned an unexpected format even after retries. Please try again or simplify the input. Raw output hint: {str(raw_output)[:500]}...")
    except Exception as e:
        logger.error(f"Error during dialogue generation: {e}")
        error_str = str(e).lower()
        if "authentication" in error_str:
             raise gr.Error("Authentication error with API. Please check your API key.")
        elif "rate limit" in error_str:
             raise gr.Error("API rate limit exceeded. Please wait and try again, or check your usage tier.")
        elif "base_url" in error_str or "connection" in error_str: # This refers to OpenAI connection for dialogue
            dialogue_base_url = resolved_openai_base_url or 'the configured OpenAI endpoint for dialogue generation'
            raise gr.Error(f"Could not connect to {dialogue_base_url}. Please check the URL and network connection.")
        elif "invalid request" in error_str and "image" in error_str:
             raise gr.Error("Error processing image with Vision API. The image might be invalid, unsupported, or the model doesn't support image input.")
        else:
            raise gr.Error(f"An error occurred during dialogue generation: {e}")

    if not llm_output or not llm_output.dialogue:
        raise gr.Error("The AI failed to generate a dialogue script. The input might be too short or unclear.")

    characters = 0
    total_lines = len(llm_output.dialogue)
    logger.info(f"Starting TTS generation for {total_lines} dialogue lines.")
    gr.Info(f"🪄 Generating audio for {total_lines} dialogue lines... (this may take a while)")

    results = [None] * total_lines

    with cf.ThreadPoolExecutor(max_workers=10) as executor:
        future_to_index = {}
        # The comprehensive API key check for resolved_openai_api_key (used by get_mp3 for one-api)
        # was done at the beginning of the generate_audio function.
        # If that check passed, we can proceed.
        future_to_index = {
            executor.submit(get_mp3, line.text, line.voice(), resolved_openai_api_key): i # English only, no language parameter needed
            for i, line in enumerate(llm_output.dialogue) if line.text.strip()
        }
        
        for line in llm_output.dialogue:
            if line.text.strip():
                characters += len(line.text)

        processed_count = 0
        tts_start_time = time.time()
        for future in cf.as_completed(future_to_index):
            index = future_to_index[future]
            line_obj = llm_output.dialogue[index]
            transcript_line = f"{line_obj.speaker}: {line_obj.text}"
            try:
                audio_chunk = future.result()
                results[index] = (transcript_line, audio_chunk)
                processed_count += 1
                gr.Info(f"🪄 Generated audio for {processed_count}/{total_lines} lines...")
            except Exception as exc:
                 logger.error(f'TTS generation failed for line {index+1} after retries: {exc}')
                 error_msg = f"[TTS Error: Failed audio for line {index+1}]"
                 results[index] = (transcript_line, error_msg) 

    logger.info(f"TTS generation took {time.time() - tts_start_time:.2f} seconds.")
    logger.info(f"Total characters for TTS: {characters}")

    gr.Info("🧩 Combining audio segments...")
    final_audio_chunks = []
    final_transcript_lines = []
    successful_lines = 0
    for i, result in enumerate(results):
        line_obj = llm_output.dialogue[i] 
        if result is None:
            if line_obj.text.strip(): 
                logger.error(f"Result missing for non-empty line {i+1}. Original text: '{line_obj.text[:50]}...' Skipping.")
                final_transcript_lines.append(f"[Internal Error: Audio result missing for line {i+1}] {line_obj.speaker}: {line_obj.text}")
            continue 

        transcript_part, audio_part = result
        final_transcript_lines.append(transcript_part) 
        if isinstance(audio_part, bytes):
            final_audio_chunks.append(audio_part)
            successful_lines += 1
    
    if not final_audio_chunks:
        if any("[TTS Error" in line for line in final_transcript_lines):
             raise gr.Error("Failed to generate audio for all lines. Please check the transcript for details and review API key/status.")
        else:
             raise gr.Error("Failed to generate any audio, although dialogue script was created. Check TTS service status or API key.")

    audio = b"".join(final_audio_chunks)
    transcript = "\n\n".join(final_transcript_lines)

    # Build HTML transcript for color-coded display
    html_transcript_lines = []
    for line in final_transcript_lines:
        if ": " in line:
            speaker, text = line.split(": ", 1)
            color = TRANSCRIPT_COLORS.get(speaker, "#000000")
            escaped_text = html.escape(text)
            html_line = f'<div class="transcript-bubble" style="background-color: {color}; padding: 10px; margin: 0; border-radius: 10px; color: black; box-shadow: 0 2px 4px rgba(0,0,0,0.1);"><strong>{speaker}:</strong> {escaped_text}</div>'
            html_transcript_lines.append(html_line)
        else:
            html_transcript_lines.append(html.escape(line))
    
    # Build learning notes HTML
    learning_notes_html = f"""
    <div class="learning-notes-container" style="margin-top: 30px; padding: 20px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.2);">
        <h2 style="color: white; text-align: center; margin-bottom: 25px; font-size: 28px; text-shadow: 2px 2px 4px rgba(0,0,0,0.3);">📚 Study Notes 學習筆記</h2>
        
        <!-- Ideas Section -->
        <div class="notes-section" style="background: white; padding: 20px; border-radius: 10px; margin-bottom: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <h3 style="color: #667eea; border-bottom: 3px solid #667eea; padding-bottom: 10px; margin-bottom: 15px;">💡 Ideas 討論要點</h3>
            <div style="line-height: 1.8; color: #333;">
                {llm_output.learning_notes.ideas}
            </div>
        </div>
        
        <!-- Language Section -->
        <div class="notes-section" style="background: white; padding: 20px; border-radius: 10px; margin-bottom: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <h3 style="color: #764ba2; border-bottom: 3px solid #764ba2; padding-bottom: 10px; margin-bottom: 15px;">📖 Language 語言學習</h3>
            <div style="overflow-x: auto;">
                {llm_output.learning_notes.language}
            </div>
        </div>
        
        <!-- Communication Strategies Section -->
        <div class="notes-section" style="background: white; padding: 20px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <h3 style="color: #f093fb; border-bottom: 3px solid #f093fb; padding-bottom: 10px; margin-bottom: 15px;">💬 Communication Strategies 溝通策略</h3>
            <div style="line-height: 1.8; color: #333;">
                {llm_output.learning_notes.communication_strategies}
            </div>
        </div>
    </div>
    """
    
    html_transcript = f'<div class="transcript-container" style="max-height: 400px; overflow-y: auto; background-color: #f9f9f9; padding: 10px; border-radius: 5px;">{"<br>".join(html_transcript_lines)}</div>{learning_notes_html}'

    logger.info(f"Successfully generated audio for {successful_lines}/{total_lines} lines.")

    temporary_directory = "./gradio_cached_files/tmp/" 
    os.makedirs(temporary_directory, exist_ok=True)

    try:
        temp_file_path = None
        with NamedTemporaryFile(
            dir=temporary_directory,
            delete=False, 
            suffix=".mp3",
            prefix="GI_audio_"
        ) as temp_file:
             temp_file.write(audio)
             temp_file_path = temp_file.name 

        if temp_file_path:
             logger.info(f"Audio saved to temporary file: {temp_file_path}")
        else:
             raise IOError("Temporary file path was not obtained.")

    except Exception as e:
        logger.error(f"Failed to write temporary audio file: {e}")
        raise gr.Error("Failed to save the generated audio file.")

    try:
        for file in glob.glob(f"{temporary_directory}GI_audio_*.mp3"):
            if os.path.isfile(file) and time.time() - os.path.getmtime(file) > 24 * 60 * 60: 
                try:
                    os.remove(file)
                    logger.debug(f"Removed old temp file: {file}")
                except OSError as e_rem:
                     logger.warning(f"Could not remove old temp file {file}: {e_rem}") 
    except Exception as e: 
        logger.warning(f"Error during old temp file cleanup: {e}")

    total_duration = time.time() - start_time
    tts_cost = (characters / 1_000_000) * 15 * 7.8
    gr.Info(f"🎉 Audio generation complete! Total time: {total_duration:.2f} seconds.")
    gr.Info(f"💸 This audio generation costs HK${tts_cost:.2f}.")

    # Prepare audio title for history
    # Get current time in UTC
    utc_now = datetime.datetime.now(datetime.timezone.utc)
    # Define Hong Kong timezone
    hk_tz = pytz.timezone('Asia/Hong_Kong')
    # Convert UTC time to Hong Kong time
    hk_now = utc_now.astimezone(hk_tz)
    # Format the Hong Kong time
    final_podcast_title = f"{podcast_title_base} - {hk_now.strftime('%Y-%m-%d %H:%M')}"
    
    # Escape HTML transcript (with learning notes) for JavaScript string literal
    # We need to escape the full HTML version so it includes learning notes when saved
    escaped_html_transcript = html_transcript.replace('\\', '\\\\').replace("'", "\\'").replace('"', '\\"').replace('\n', '\\n').replace('\r', '\\r')

    # Create JavaScript to call the save function in head.html
    # The audio file path needs to be accessible by the client's browser.
    # Gradio serves files from a temporary location. We need to ensure this path is correct.
    # If temp_file_path is absolute, we might need to make it relative or ensure it's served.
    # For Gradio, files in `file_output` are typically served under `/file=...`
    # We assume `temp_file_path` as returned by `gr.Audio` is directly fetchable.
    
    # The audio_output component in Gradio will have a URL like /file=/path/to/temp_file.mp3
    # We need to pass this server-relative path to the JS function.
    # `temp_file_path` is the absolute path on the server.
    # Gradio's `gr.Audio(type="filepath")` returns the absolute path.
    # When this path is set as the value of an `gr.Audio` output, Gradio makes it accessible via a `/file=` URL.
    # The JS in `head.html` will fetch this URL.

    # Construct the web-accessible URL for the temporary file
    # Gradio serves files from gr.Audio(type="filepath") via /file=<path>
    # Ensure the path is properly URL encoded if it contains special characters, though Gradio might handle this.
    # For simplicity, we'll assume basic paths for now. If issues persist, URL encoding might be needed.
    # gradio_file_url = f"/file={temp_file_path}" # Original absolute path approach

    temp_file_path_obj = Path(temp_file_path)
    try:
        # Path.cwd() in the context of the uvicorn server (likely /app in Docker)
        app_root = Path.cwd()
        relative_temp_path = temp_file_path_obj.relative_to(app_root)
        gradio_file_url = f"/file={relative_temp_path}"
        logger.info(f"Constructed relative gradio_file_url: {gradio_file_url}")
    except ValueError:
        # Fallback if relative_to fails (e.g. different drives or not a subpath)
        gradio_file_url = f"/file={temp_file_path}" # Keep the absolute path
        logger.warning(f"Could not make path relative to CWD ('{app_root}'). Using absolute path for gradio_file_url: {gradio_file_url}")


    # The audio_url will now be fetched by JS from a hidden gr.File component
    # We pass the raw temp_file_path to that component.
    # The JSON will just contain a marker or the title to correlate.
    data_to_send = {
        "title": final_podcast_title,
        # "audio_url": gradio_file_url, # REMOVED - JS will get this from hidden gr.File
        "audio_file_component_id": "temp_audio_file_url_holder", # ID of the hidden gr.File
        "transcript": escaped_html_transcript,  # Now includes learning notes
        "tts_cost": f"{tts_cost:.2f}" # Added tts_cost, formatted as string
    }
    json_data_string = json.dumps(data_to_send)
    
    logger.debug(f"Returning JSON data for JS trigger (no audio_url, JS will fetch from component): {json_data_string[:200]}...")

    return temp_file_path, html_transcript, json_data_string, temp_file_path # 4th item for hidden gr.File


def generate_word_document(transcript_html: str, title: str = "Group Discussion Notes") -> str:
    """
    Generates a Word document from the transcript and study notes HTML.
    Returns the path to the generated Word document.
    """
    try:
        import re
        from datetime import datetime
        from bs4 import BeautifulSoup
        
        # Create a new Word document
        doc = docx.Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add timestamp in Hong Kong time
        hk_tz = pytz.timezone('Asia/Hong_Kong')
        hk_now = datetime.now(hk_tz)
        timestamp = hk_now.strftime("%Y-%m-%d %H:%M:%S")
        timestamp_para = doc.add_paragraph(f"Generated by Mr.🆖 DiscussAI on {timestamp}")
        timestamp_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # Add a blank line
        
        # Parse the HTML to extract transcript and learning notes
        soup = BeautifulSoup(transcript_html, 'html.parser')
        
        # Debug logging
        logger.debug(f"Parsing HTML for Word document. HTML length: {len(transcript_html)}")
        logger.debug(f"Found transcript containers: {len(soup.find_all('div', class_='transcript-container'))}")
        logger.debug(f"Found transcript bubbles: {len(soup.find_all('div', class_='transcript-bubble'))}")
        
        # Find transcript container
        transcript_container = soup.find('div', class_='transcript-container')
        # Get transcript bubbles - either from container or directly from soup
        transcript_bubbles = []
        if transcript_container:
            transcript_bubbles = transcript_container.find_all('div', class_='transcript-bubble')
        else:
            # Fallback: try to find bubbles directly (in case HTML structure is different)
            transcript_bubbles = soup.find_all('div', class_='transcript-bubble')
        
        if transcript_bubbles:
            # Add transcript heading
            doc.add_heading("Transcript 文字稿", level=1)
            
            logger.debug(f"Processing {len(transcript_bubbles)} transcript bubbles")
            
            # Process transcript bubbles
            for idx, bubble in enumerate(transcript_bubbles):
                # Extract speaker and text
                bubble_text = bubble.get_text()
                
                # Skip if this bubble contains the entire transcript (duplication issue)
                # This happens when the HTML structure is malformed
                if len(bubble_text) > 5000 or bubble_text.count('Candidate A:') > 2:
                    logger.warning(f"Skipping bubble {idx} - appears to contain entire transcript (length: {len(bubble_text)})")
                    continue
                
                if ': ' in bubble_text:
                    speaker, text = bubble_text.split(': ', 1)
                    
                    # Validate speaker name
                    if speaker not in TRANSCRIPT_COLORS:
                        logger.warning(f"Skipping bubble {idx} - invalid speaker: {speaker}")
                        continue
                    
                    # Add speaker paragraph with formatting
                    speaker_para = doc.add_paragraph()
                    speaker_run = speaker_para.add_run(f"{speaker}:")
                    speaker_run.bold = True
                    speaker_run.font.size = Pt(11)
                    
                    # Get speaker color from TRANSCRIPT_COLORS
                    speaker_color = TRANSCRIPT_COLORS.get(speaker, "#000000")
                    # Convert hex to RGB
                    hex_color = speaker_color.lstrip('#')
                    rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                    
                    # Add a colored background paragraph for the text
                    text_para = doc.add_paragraph()
                    text_para.paragraph_format.left_indent = Inches(0.25)
                    text_run = text_para.add_run(text)
                    text_run.font.size = Pt(11)
                    
                    # Add shading to the paragraph
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), hex_color)
                    text_para._p.get_or_add_pPr().append(shading_elm)
                    
                    doc.add_paragraph()  # Add spacing between bubbles
        
        # Find learning notes container
        learning_notes_container = soup.find('div', class_='learning-notes-container')
        if learning_notes_container:
            doc.add_page_break()  # Add a page break before learning notes
            
            # Add learning notes heading
            doc.add_heading("Study Notes 學習筆記", level=1)
            
            # Process each section
            sections = learning_notes_container.find_all('div', class_='notes-section')
            
            for section in sections:
                # Get section heading
                section_heading = section.find('h3')
                heading_text = section_heading.get_text() if section_heading else ""
                if section_heading:
                    doc.add_heading(heading_text, level=2)
                
                # Process section content
                section_content = section.find('div')
                if section_content:
                    # Handle tables (for language section)
                    tables = section_content.find_all('table')
                    if tables:
                        for table in tables:
                            # Create Word table
                            rows = table.find_all('tr')
                            if rows and rows[0].find_all(['th', 'td']):
                                word_table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
                                word_table.style = 'Table Grid'
                                word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                
                                for i, row in enumerate(rows):
                                    cells = row.find_all(['th', 'td'])
                                    for j, cell in enumerate(cells):
                                        word_table.cell(i, j).text = cell.get_text().strip()
                                        
                                        # Make header cells bold
                                        if cell.name == 'th' and word_table.cell(i, j).paragraphs[0].runs:
                                            word_table.cell(i, j).paragraphs[0].runs[0].bold = True
                    
                    # Handle Communication Strategies section
                    elif "Communication Strategies" in heading_text or "溝通策略" in heading_text:
                        # Parse HTML content to extract strategies
                        html_str = str(section_content)
                        
                        # Split by <strong> tags to find strategy titles
                        strategy_parts = re.split(r'<strong>(\d+\.\s+[^<]+)</strong>', html_str)
                        
                        for i in range(1, len(strategy_parts), 2):
                            if i < len(strategy_parts):
                                # Strategy title - remove the number prefix since it's already in numbered list
                                strategy_title = strategy_parts[i].strip()
                                # Remove the number and dot at the beginning (e.g., "1. " -> "")
                                strategy_title_clean = re.sub(r'^\d+\.\s+', '', strategy_title)
                                para = doc.add_paragraph(strategy_title_clean, style='List Number')
                                if para.runs:
                                    para.runs[0].bold = True
                                
                                # Strategy content (examples and explanations)
                                if i + 1 < len(strategy_parts):
                                    content = strategy_parts[i + 1]
                                    
                                    # Extract examples - try multiple patterns for robustness
                                    examples = []
                                    
                                    # Pattern 1: Text in <em> tags with quotes
                                    examples.extend(re.findall(r'<em>[\u0022\u201C\u201D]([^\u0022\u201C\u201D]+)[\u0022\u201C\u201D]</em>', content))
                                    
                                    # Pattern 2: Text in <em> tags without quotes
                                    if not examples:
                                        examples.extend(re.findall(r'<em>([^<]+)</em>', content))
                                    
                                    # Pattern 3: Lines starting with bullet point (•) followed by quoted text
                                    if not examples:
                                        examples.extend(re.findall(r'•\s*[\u0022\u201C\u201D]([^\u0022\u201C\u201D]+)[\u0022\u201C\u201D]', content))
                                    
                                    # Pattern 4: Any quoted text in the content
                                    if not examples:
                                        examples.extend(re.findall(r'[\u0022\u201C\u201D]([^\u0022\u201C\u201D]+)[\u0022\u201C\u201D]', content))
                                    
                                    # Add examples as indented bullet points
                                    for example in examples:
                                        example = example.strip()
                                        if example:  # Only add non-empty examples
                                            # Add quotes if not already present
                                            if not (example.startswith('"') or
                                            example.startswith('\u201C')):
                                                example = f'"{example}"'
                                            example_para = doc.add_paragraph(example, style='List Bullet')
                                            example_para.paragraph_format.left_indent = Inches(0.5)
                                    
                                    # Extract ALL Chinese text (not just those starting with '用於')
                                    # Remove HTML tags first
                                    clean_content = re.sub(r'<[^>]+>', '', content)
                                    # Find Chinese sentences (characters followed by punctuation)
                                    chinese_sentences = re.findall(r'[\u4e00-\u9fff]+[^<\n]*[。！？]?', clean_content)
                                    
                                    # Add Chinese explanations as indented bullet points
                                    for chinese in chinese_sentences:
                                        chinese = chinese.strip()
                                        # Skip if it's empty or just punctuation
                                        if chinese and not re.match(r'^[。！？，、；：""''（）]+$', chinese):
                                            chinese_para = doc.add_paragraph(chinese, style='List Bullet')
                                            chinese_para.paragraph_format.left_indent = Inches(0.5)
                    
                    # Handle Ideas section
                    elif "Ideas" in heading_text or "討論要點" in heading_text:
                        # Parse HTML content to extract questions and points
                        html_str = str(section_content)
                        
                        # Split by <strong> tags to find questions
                        question_parts = re.split(r'<strong>([^<]+)</strong>', html_str)
                        
                        for i in range(1, len(question_parts), 2):
                            if i < len(question_parts):
                                # Add spacing before each question (except the first one)
                                if i > 1:
                                    doc.add_paragraph()  # Add blank line for spacing
                                
                                # Question title
                                question_title = question_parts[i].strip()
                                para = doc.add_paragraph()
                                run = para.add_run(question_title)
                                run.bold = True
                                
                                # Question content (main points and sub-points)
                                if i + 1 < len(question_parts):
                                    content = question_parts[i + 1]
                                    
                                    # Split by <br> to get individual lines
                                    lines = re.split(r'<br\s*/?>', content)
                                    
                                    for line in lines:
                                        # Clean HTML tags
                                        clean_line = re.sub(r'<[^>]+>', '', line).strip()
                                        clean_line = clean_line.replace('&nbsp;', ' ')
                                        
                                        if not clean_line:
                                            continue
                                        
                                        # Check if it's a main point (starts with •)
                                        if clean_line.startswith('•'):
                                            # Remove the bullet and add as regular paragraph
                                            main_point = clean_line[1:].strip()
                                            doc.add_paragraph(main_point)
                                        
                                        # Check if it's a sub-point (starts with -)
                                        elif clean_line.strip().startswith('-'):
                                            # Remove the dash and add as indented bullet
                                            sub_point = clean_line.strip()[1:].strip()
                                            sub_para = doc.add_paragraph(sub_point, style='List Bullet')
                                            sub_para.paragraph_format.left_indent = Inches(0.5)
                                        
                                        # Check if line has multiple spaces (indicating indentation)
                                        elif '    ' in clean_line or clean_line.startswith('  '):
                                            # This is likely a sub-point
                                            sub_point = clean_line.strip()
                                            if sub_point.startswith('-'):
                                                sub_point = sub_point[1:].strip()
                                            sub_para = doc.add_paragraph(sub_point, style='List Bullet')
                                            sub_para.paragraph_format.left_indent = Inches(0.5)
                                        
                                        # Otherwise, it's a regular paragraph
                                        elif clean_line:
                                            doc.add_paragraph(clean_line)
        
        # Save the document
        temporary_directory = "./gradio_cached_files/tmp/"
        os.makedirs(temporary_directory, exist_ok=True)
        
        # Use NamedTemporaryFile for consistent naming with audio files
        with NamedTemporaryFile(
            dir=temporary_directory,
            delete=False,
            suffix=".docx",
            prefix="GI_notes_"
        ) as temp_doc_file:
            doc.save(temp_doc_file.name)
            doc_path = temp_doc_file.name
        
        # Clean up old Word documents (older than 60 minutes)
        try:
            for file in glob.glob(f"{temporary_directory}GI_notes_*.docx"):
                if os.path.isfile(file) and time.time() - os.path.getmtime(file) > 60 * 60:  # 60 minutes
                    try:
                        os.remove(file)
                        logger.debug(f"Removed old Word document: {file}")
                    except OSError as e_rem:
                        logger.warning(f"Could not remove old Word document {file}: {e_rem}")
        except Exception as e:
            logger.warning(f"Error during old Word document cleanup: {e}")
        
        logger.info(f"Word document generated successfully: {doc_path}")
        return doc_path
        
    except Exception as e:
        logger.error(f"Error generating Word document: {e}")
        raise gr.Error(f"Failed to generate Word document: {e}")


# --- Gradio UI Definition ---

allowed_extensions = [
    ".jpg", ".jpeg", ".png", ".docx", ".pdf"
]

examples_dir = Path("examples")
examples = [
    [ # Input method, dialogue mode, files, text, api_key
        "Upload Files", [str(examples_dir / "DSE 2023 Paper 4 Set 1.2.png")], "", "Deeper", None
    ],
    [
        "Upload Files", [str(examples_dir / "DSE 2024 Paper 4 Set 5.2.png")], "", "Normal",  None
    ]
]

def read_file_content(filepath: str, default: str = "") -> str:
    try:
        return Path(filepath).read_text(encoding='utf-8') 
    except FileNotFoundError:
        logger.warning(f"{filepath} not found, using default content.")
        return default
    except Exception as e:
         logger.error(f"Error reading file {filepath}: {e}. Using default.")
         return default


description_md = read_file_content("description.md", "AI-Powered Group Discussion Practice for HKDSE Oral Exam")
footer_md = read_file_content("footer.md", "")
head_html = read_file_content("head.html", "")


with gr.Blocks(theme="ocean", title="Mr.🆖 DiscussAI 👥🎙️", css="footer{display:none !important}") as demo: # Reverted allowed_paths
    gr.Markdown(description_md)

    with gr.Row():
        input_method_radio = gr.Radio(
            ["Upload Files", "Enter Topic"],
            label="📁 Discussion Topic",
            value="Upload Files"
        )

    with gr.Group(visible=True) as file_upload_group:
        file_input = gr.Files(
            label="📸 Upload Group Interaction Task",
            file_types=allowed_extensions,
            file_count="multiple",
        )

    with gr.Group(visible=False) as text_input_group: 
        text_input = gr.Textbox(
            label="✍️ Enter Topic",
            lines=10,
            placeholder="Paste or type your discussion topic here..."
        )

    with gr.Row():
        dialogue_mode_radio = gr.Radio(
            ["Normal", "Deeper"],
            label="🎯 Depth of Discussion",
            value="Normal",
            info="Select 'Deeper' if you prefer a more detailed discussion with further ideas and elaborations"
        )

    API_KEY_URL = "https://api.mr5ai.com"
    with gr.Accordion("⚙️ Advanced Settings", open=False):
        gr.Markdown(
            f"💡 Get your Mr.🆖 AI Hub API Key [here]({API_KEY_URL})"
        )
        api_key_input = gr.Textbox(
                label="Mr.🆖 AI Hub API Key",
                type="password",
                placeholder="sk-xxx",
                elem_id="mr_ng_ai_hub_api_key_input"
        )

    submit_button = gr.Button("✨ Generate Discussion with Audio and Study Notes", variant="primary")

    audio_output = gr.Audio(label="🔊 Audio", type="filepath", elem_id="podcast_audio_player") # Keep existing elem_id
    transcript_output = gr.HTML(label="📃 Transcript", elem_id="podcast_transcript_display") # Keep existing elem_id
    
    # Hidden textbox to store transcript HTML for Word document generation
    # This will be updated by JavaScript when loading from archives
    transcript_html_storage = gr.Textbox(label="Transcript HTML Storage", visible=False, elem_id="transcript_html_storage")
    
    # Add download button for Word document
    with gr.Row():
        download_word_btn = gr.Button("📄 Download as Word Document →", variant="primary")
        word_doc_output = gr.File(label="📄 Word Document", visible=True)

    with gr.Accordion("📜 Archives (Stored in your browser)", open=False): # Keep existing Accordion
        # This HTML component will be populated by JavaScript from head.html
        podcast_history_display = gr.HTML("<ul id='podcastHistoryList' style='list-style-type: none; padding: 0;'><li>Loading archives...</li></ul>")
        # Hidden Textbox component to pass JSON data to JavaScript
        js_trigger_data_textbox = gr.Textbox(label="JS Trigger Data", visible=False, elem_id="js_trigger_data_textbox")
        # Hidden File component to get a Gradio-served URL for the audio
        temp_audio_file_output_for_url = gr.File(label="Temp Audio File URL Holder", visible=False, elem_id="temp_audio_file_url_holder")


    def switch_input_method(choice):
        """Updates visibility and clears the inactive input fields."""
        is_upload = choice == "Upload Files"
        is_text = choice == "Enter Topic"

        # Determine visibility updates
        file_vis = is_upload
        text_vis = is_text

        # Determine value updates (clear hidden fields)
        # gr.update() means no change to value
        file_val_update = gr.update(value=None) if not is_upload else gr.update()
        text_val_update = gr.update(value="") if not is_text else gr.update()

        return {
            file_upload_group: gr.update(visible=file_vis),
            text_input_group: gr.update(visible=text_vis),
            file_input: file_val_update,
            text_input: text_val_update,
        }

    input_method_radio.change(
        fn=switch_input_method,
        inputs=input_method_radio,
        outputs=[
            file_upload_group, 
            text_input_group, 
            file_input, 
            text_input
        ]
    )

    # Function to sync transcript to hidden storage
    def sync_transcript_to_storage(transcript_html):
        """Syncs the transcript HTML to the hidden storage for Word download."""
        return transcript_html
    
    submit_button.click(
        fn=generate_audio,
        inputs=[ # Order must match generate_audio parameters
            input_method_radio,
            file_input,
            text_input,
            dialogue_mode_radio,
            api_key_input
        ],
        outputs=[audio_output, transcript_output, js_trigger_data_textbox, temp_audio_file_output_for_url],
        api_name="generate_audio"
    ).then(
        fn=sync_transcript_to_storage,
        inputs=[transcript_output],
        outputs=[transcript_html_storage]
    )
    
    # Function to handle Word document download
    def handle_word_download(transcript_html_from_storage):
        """
        Generates a Word document from the transcript HTML.
        Uses the hidden storage textbox which is updated both when generating new discussions
        and when loading from archives (via JavaScript).
        """
        if not transcript_html_from_storage or not transcript_html_from_storage.strip():
            raise gr.Error("No transcript available to download. Please generate a discussion first or load one from Archives.")
        
        try:
            # Extract title from the HTML or use a default
            title = "Group Discussion Notes"
            
            # Try to extract title from the HTML if it contains timestamp
            # Look for patterns like "filename - YYYY-MM-DD HH:MM" in the transcript
            import re
            from bs4 import BeautifulSoup
            
            # The transcript HTML from storage might have escaped characters from JavaScript
            # Unescape them before processing
            transcript_html_unescaped = transcript_html_from_storage.replace('\\n', '\n').replace('\\r', '\r').replace('\\"', '"').replace("\\'", "'").replace('\\\\', '\\')
            
            # Parse HTML to check if it's valid
            soup = BeautifulSoup(transcript_html_unescaped, 'html.parser')
            
            # Check if we have actual content (transcript bubbles or learning notes)
            has_content = bool(soup.find('div', class_='transcript-bubble') or
                             soup.find('div', class_='learning-notes-container'))
            
            if not has_content:
                raise gr.Error("The transcript appears to be empty or invalid. Please generate a discussion first or load one from Archives.")
            
            # Generate the Word document
            doc_path = generate_word_document(transcript_html_unescaped, title)
            
            # Return the file path for Gradio's File component
            return doc_path
        except gr.Error:
            # Re-raise Gradio errors as-is
            raise
        except Exception as e:
            logger.error(f"Error in handle_word_download: {e}")
            raise gr.Error(f"Failed to generate Word document: {str(e)}")
    
    # Connect the download button to the function
    download_word_btn.click(
        fn=handle_word_download,
        inputs=[transcript_html_storage],
        outputs=[word_doc_output],
        api_name="download_word"
    )

    # Create a wrapper function that also syncs transcript to storage for examples
    def generate_audio_with_sync(input_method, files, input_text, dialogue_mode, openai_api_key):
        """Wrapper that generates audio and syncs transcript to storage."""
        audio_path, transcript_html, json_data, temp_audio = generate_audio(
            input_method, files, input_text, dialogue_mode, openai_api_key
        )
        # Return all outputs including the synced transcript storage
        return audio_path, transcript_html, json_data, temp_audio, transcript_html
    
    gr.Examples(
        examples=examples,
        inputs=[ # Ensure order matches generate_audio parameters for examples
            input_method_radio,
            file_input,
            text_input,
            dialogue_mode_radio,
            api_key_input
        ],
        # Now includes transcript_html_storage in outputs so Word download works for examples
        outputs=[audio_output, transcript_output, js_trigger_data_textbox, temp_audio_file_output_for_url, transcript_html_storage],
        fn=generate_audio_with_sync,
        cache_examples=True,
        run_on_click=True,
        label="Examples (Click for Demo)"
    )

    gr.Markdown(footer_md)
    demo.head = (os.getenv("HEAD", "") or "") + head_html

# --- App Setup & Launch ---

demo = demo.queue(
    max_size=20,
    default_concurrency_limit=5, 
)

app = gr.mount_gradio_app(app, demo, path="/")

if __name__ == "__main__":
    examples_dir.mkdir(exist_ok=True)
    example_files = [
        "DSE 2023 Paper 4 Set 1.2.png",
        "DSE 2024 Paper 4 Set 5.2.png"
    ]
    for fname in example_files:
        fpath = examples_dir / fname
        if not fpath.is_file():
            logger.warning(f"Example file {fpath} not found. Creating empty placeholder.")
            try:
                fpath.touch()
            except OSError as e:
                logger.error(f"Failed to create placeholder file {fpath}: {e}")

    os.makedirs("./gradio_cached_files/tmp/", exist_ok=True)

    logger.info("Starting Gradio application via Uvicorn...")
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
